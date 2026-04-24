"""python-docx helpers for the per-persona KT bundle (Phase 2).

Produces a single `.docx` per persona with Calibri body, navy H1/H2, grey
code-block shading, a running header (`{tenant} — KT for {persona}`) and a
page-number footer. Image insertion auto-converts any size to a fixed
printable width so the document stays inside the page margins.

Safe to import even when `python-docx` is missing — the ImportError is
deferred until a caller actually instantiates `DocxBuilder`, which lets
flag-off builds skip the dependency entirely.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Iterable

logger = logging.getLogger(__name__)


# Style colours. Navy palette picked to match the legacy Markdown/PDF
# output so the bundle looks like a natural extension of the existing
# CrawlDocGenerator report, not a separate product.
_H1_COLOR = (0x0F, 0x6F, 0xC6)  # navy
_H2_COLOR = (0x1F, 0x4E, 0x79)  # dark navy
_H3_COLOR = (0x32, 0x5C, 0x8A)
_CODE_SHADING = "F2F2F2"
_CALLOUT_DEFAULT = "0F6FC6"


class DocxBuilder:
    """Minimal wrapper around python-docx with Carbon-ish styling."""

    def __init__(self, title: str, subtitle: str, tenant: str, persona: str) -> None:
        try:
            from docx import Document  # noqa: F401 — imported lazily
            from docx.shared import Pt, RGBColor, Inches, Cm  # noqa: F401
        except ImportError as e:  # pragma: no cover — dep check
            raise ImportError(
                "python-docx is required for DOCX generation. "
                "Install with `pip install python-docx Pillow` or flip "
                "`kt_pro.docx_bundle.enabled = false` in config.yaml."
            ) from e

        from docx import Document
        from docx.shared import Pt

        self._Document = Document
        self._Pt = Pt
        self.title = title
        self.subtitle = subtitle
        self.tenant = tenant
        self.persona = persona
        self.doc = Document()

        self._configure_base_style()
        self._configure_header_footer()
        self._render_title_page()

    # --- Style setup -----------------------------------------------------
    def _configure_base_style(self) -> None:
        """Calibri 11pt body + navy Heading 1/2 to match the brand palette."""
        from docx.shared import Pt, RGBColor

        # Body (Normal) style
        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

        def _color(name: str, rgb: tuple[int, int, int], size_pt: int) -> None:
            try:
                style = styles[name]
            except KeyError:
                return
            style.font.name = "Calibri"
            style.font.size = Pt(size_pt)
            style.font.color.rgb = RGBColor(*rgb)
            style.font.bold = True

        _color("Heading 1", _H1_COLOR, 20)
        _color("Heading 2", _H2_COLOR, 15)
        _color("Heading 3", _H3_COLOR, 12)

    def _configure_header_footer(self) -> None:
        """Running header '{tenant} — KT for {persona}', page X footer."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        section = self.doc.sections[0]
        header = section.header.paragraphs[0]
        header.text = f"{self.tenant} — KT for {self.persona}"

        footer = section.footer.paragraphs[0]
        run = footer.add_run("Page ")

        def _field(code: str) -> None:
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText")
            instr.text = code
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            footer_run = footer.add_run()
            footer_run._r.append(fld_begin)
            footer_run._r.append(instr)
            footer_run._r.append(fld_end)

        _field(" PAGE ")
        footer.add_run(" of ")
        _field(" NUMPAGES ")

    def _render_title_page(self) -> None:
        """Document heading. Intentionally simple — no cover art."""
        from docx.shared import Pt, RGBColor

        title_p = self.doc.add_paragraph()
        run = title_p.add_run(self.title)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*_H1_COLOR)

        if self.subtitle:
            sub_p = self.doc.add_paragraph()
            sub_run = sub_p.add_run(self.subtitle)
            sub_run.font.size = Pt(13)
            sub_run.font.color.rgb = RGBColor(*_H2_COLOR)
            sub_run.italic = True

        meta_p = self.doc.add_paragraph()
        meta_run = meta_p.add_run(
            f"Tenant: {self.tenant}    |    Persona: {self.persona}"
        )
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(0x52, 0x52, 0x52)

        self.doc.add_paragraph()  # spacer

    # --- Public API ------------------------------------------------------
    def h1(self, text: str) -> None:
        self.doc.add_heading(text, level=1)

    def h2(self, text: str) -> None:
        self.doc.add_heading(text, level=2)

    def h3(self, text: str) -> None:
        self.doc.add_heading(text, level=3)

    def p(self, text: str) -> None:
        if text:
            self.doc.add_paragraph(text)

    def bullet(self, text: str) -> None:
        self.doc.add_paragraph(text, style="List Bullet")

    def numbered(self, text: str) -> None:
        self.doc.add_paragraph(text, style="List Number")

    def code_block(self, text: str) -> None:
        """Shaded Consolas 10pt block. Preserves newlines."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt

        para = self.doc.add_paragraph()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), _CODE_SHADING)
        pPr = para._p.get_or_add_pPr()
        pPr.append(shading)
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    def callout(self, label: str, text: str, color: str = _CALLOUT_DEFAULT) -> None:
        """Label/message one-liner with a shaded background.

        `color` is a hex RGB string without leading `#`.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt, RGBColor

        para = self.doc.add_paragraph()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), color)
        pPr = para._p.get_or_add_pPr()
        pPr.append(shading)

        label_run = para.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        label_run.font.size = Pt(11)

        body_run = para.add_run(text)
        body_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        body_run.font.size = Pt(11)

    def table(
        self,
        header: list[str],
        rows: list[list[str]],
        widths: list[int] | None = None,
    ) -> None:
        """Insert a simple table with a header row in navy.

        `widths` — optional column widths in inches. Ignored when shorter
        than `header`; python-docx will use its default.
        """
        from docx.shared import Inches, RGBColor

        if not header:
            return
        t = self.doc.add_table(rows=1, cols=len(header))
        t.style = "Light Grid Accent 1"

        hdr_cells = t.rows[0].cells
        for i, label in enumerate(header):
            hdr_cells[i].text = str(label)
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(*_H2_COLOR)

        for row in rows:
            row_cells = t.add_row().cells
            for i, val in enumerate(row):
                if i < len(row_cells):
                    row_cells[i].text = "" if val is None else str(val)

        if widths and len(widths) == len(header):
            for row in t.rows:
                for idx, width_in in enumerate(widths):
                    try:
                        row.cells[idx].width = Inches(width_in)
                    except Exception:
                        pass

    def figure(self, png_path: str, caption: str, *, width_px: int = 620) -> None:
        """Insert an image with caption. Silently skips if file missing."""
        from docx.shared import Inches, Pt, RGBColor

        p = Path(png_path)
        if not p.exists():
            logger.warning("Figure missing — skipping: %s", png_path)
            return
        try:
            # 96 dpi baseline — 620px ≈ 6.46" which fits the default page.
            width_in = width_px / 96.0
            self.doc.add_picture(str(p), width=Inches(width_in))
        except Exception as e:
            logger.warning("Figure insert failed for %s: %s", png_path, e)
            return
        if caption:
            cap = self.doc.add_paragraph()
            cap_run = cap.add_run(f"Figure: {caption}")
            cap_run.italic = True
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGBColor(0x52, 0x52, 0x52)

    def citations(self, refs: Iterable[str]) -> None:
        """Render a references appendix. Each ref is already formatted text."""
        self.h2("References")
        for ref in refs:
            self.bullet(ref)

    def save(self, out_path: str) -> str:
        out = Path(out_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(out))
        return str(out)
