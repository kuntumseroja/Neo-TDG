"""End-to-end composition tests using an in-memory CrawlReport + fake RAG.

The real RAGQueryEngine needs an LLM and a vector store — neither is
available in a unit test. Instead we feed `compose_all` a minimal
`CrawlReport` plus a duck-typed RAG engine that returns canned answers.
This checks that:

* Six DOCX are produced (one per persona).
* Each DOCX round-trips through `python-docx`.
* Computed sections (tech stack, endpoints, data model) pull from the
  CrawlReport.
* RAG sections call the engine with the matching persona id.
* Refused RAG responses render a visible "Evidence gap" callout.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import List

import pytest

pytest.importorskip("docx")

from docx import Document  # noqa: E402

from src.crawler.persona_composer import compose_all  # noqa: E402
from src.models.crawler import (  # noqa: E402
    CrawlReport, ProjectInfo, PackageRef, EndpointInfo, DataModel,
)
from src.rag.personas import ALL_PERSONA_IDS  # noqa: E402


# --- Stubs ----------------------------------------------------------------

@dataclass
class _FakeResponse:
    answer: str
    sources: list = field(default_factory=list)
    refused: bool = False


class _FakeRAG:
    def __init__(self, *, refuse_personas: tuple[str, ...] = ()):
        self.calls: List[dict] = []
        self.refuse_personas = refuse_personas

    def query(self, question: str, persona: str = None, **kwargs):
        self.calls.append({"q": question, "persona": persona})
        if persona in self.refuse_personas:
            return _FakeResponse(
                answer="I don't have grounded evidence for this.",
                refused=True,
            )
        return _FakeResponse(
            answer=(
                f"## Answer for {persona}\n\n"
                f"- first point [src/Foo.cs:L10-L20]\n"
                f"- second point [doc §2.1]\n\n"
                f"This paragraph cites a file [src/Bar.cs:L5]."
            ),
        )


def _sample_report() -> CrawlReport:
    return CrawlReport(
        solution="Sample.sln",
        projects=[
            ProjectInfo(
                name="Sample.Domain", path="Sample.Domain.csproj",
                layer="Domain", framework="net8.0",
                nuget_packages=[PackageRef(name="MediatR", version="12.0.0")],
            ),
            ProjectInfo(
                name="Sample.Api", path="Sample.Api.csproj",
                layer="Presentation", framework="net8.0",
                references=["Sample.Domain"],
            ),
        ],
        endpoints=[
            EndpointInfo(
                route="api/taxpayers", method="GET",
                controller="TaxpayersController",
                file="src/TaxpayersController.cs", line=15,
            ),
        ],
        data_models=[
            DataModel(
                name="Taxpayer", db_context="TaxContext",
                properties=["string Name"], relationships=[],
                file="src/Taxpayer.cs",
            ),
        ],
    )


# --- Tests ----------------------------------------------------------------

def test_compose_all_produces_six_docx(tmp_path: Path):
    produced = compose_all(
        report=_sample_report(),
        validation=None,
        tenant="Sample",
        out_dir=str(tmp_path),
        rag_engine=_FakeRAG(),
        render_pdf=False,  # keep hermetic — no soffice call
    )
    assert len(produced) == 6
    docx_files = [p for p in produced if p.suffix == ".docx"]
    assert len(docx_files) == 6
    personas_seen = {p.stem.split("_")[1] for p in docx_files}
    assert personas_seen == {"architect", "developer", "tester", "l1", "l2", "l3"}

    # Architect/developer render computed tables (projects, tech stack,
    # endpoints, data model). Tester's only computed section is "fixtures"
    # which is empty when the crawl has no test projects; L1/L2/L3 are
    # narrative-only.
    personas_with_computed = {"architect", "developer"}
    for path in docx_files:
        assert path.exists() and path.stat().st_size > 2000
        doc = Document(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Sample" in all_text
        persona_id = path.stem.split("_")[1]
        if persona_id in personas_with_computed:
            # These personas' outlines include computed sections that
            # render as tables — L1/L2 are narrative-only.
            assert len(doc.tables) >= 1, f"{persona_id} missing table"


def test_rag_sections_call_engine_with_matching_persona(tmp_path: Path):
    rag = _FakeRAG()
    compose_all(
        report=_sample_report(),
        validation=None,
        tenant="Sample",
        out_dir=str(tmp_path),
        rag_engine=rag,
        personas=["architect"],
        render_pdf=False,
    )
    # Every recorded call used the architect persona.
    assert rag.calls
    assert all(c["persona"] == "architect" for c in rag.calls)


def test_refused_rag_surfaces_callout_in_docx(tmp_path: Path):
    rag = _FakeRAG(refuse_personas=("l1",))
    produced = compose_all(
        report=_sample_report(),
        validation=None,
        tenant="Sample",
        out_dir=str(tmp_path),
        rag_engine=rag,
        personas=["l1"],
        render_pdf=False,
    )
    assert len(produced) == 1
    doc = Document(str(produced[0]))
    flat = "\n".join(p.text for p in doc.paragraphs)
    assert "Evidence gap" in flat


def test_compose_subset_personas(tmp_path: Path):
    produced = compose_all(
        report=_sample_report(),
        validation=None,
        tenant="Sample",
        out_dir=str(tmp_path),
        rag_engine=_FakeRAG(),
        personas=["developer", "tester"],
        render_pdf=False,
    )
    assert len(produced) == 2
    assert {p.stem.split("_")[1] for p in produced} == {"developer", "tester"}


def test_compose_without_rag_still_produces_docx(tmp_path: Path):
    produced = compose_all(
        report=_sample_report(),
        validation=None,
        tenant="Sample",
        out_dir=str(tmp_path),
        rag_engine=None,
        personas=["architect"],
        render_pdf=False,
    )
    assert len(produced) == 1
    assert produced[0].exists() and produced[0].stat().st_size > 1500
