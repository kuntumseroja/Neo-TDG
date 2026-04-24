"""Smoke tests for src.crawler.docx_builder.

These don't validate styling — we can't diff rendered PDFs without
LibreOffice — but they do prove:

* The file is emitted and `python-docx` can round-trip it.
* Headings, paragraphs, tables, bullets, and code blocks all contribute
  to the document body (i.e. no helper silently no-ops).
* The running header contains `tenant` + persona.
"""
from __future__ import annotations

from pathlib import Path

import pytest

pytest.importorskip("docx")

from docx import Document  # noqa: E402

from src.crawler.docx_builder import DocxBuilder  # noqa: E402


def test_docx_builder_round_trips(tmp_path: Path):
    b = DocxBuilder(
        title="CoreTax KT",
        subtitle="For Solution Architect",
        tenant="CoreTax",
        persona="Solution Architect",
    )
    b.h1("Top")
    b.h2("Section A")
    b.h3("Subsection")
    b.p("A grounded sentence [src/Foo.cs:L10-L20].")
    b.bullet("one")
    b.bullet("two")
    b.numbered("first")
    b.code_block("public void X() {}")
    b.callout("NOTE", "Watch this constraint.")
    b.table(
        header=["Name", "Layer"],
        rows=[["A", "Domain"], ["B", "Application"]],
        widths=[2, 2],
    )
    b.citations(["src/Foo.cs:L10-L20", "doc §3.2"])

    out = b.save(str(tmp_path / "out.docx"))
    assert Path(out).exists() and Path(out).stat().st_size > 2000

    doc = Document(out)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    # Title page
    assert "CoreTax KT" in all_text
    assert "Solution Architect" in all_text
    # Content
    assert "Section A" in all_text
    assert "Subsection" in all_text
    assert "one" in all_text and "two" in all_text
    assert "first" in all_text
    assert "public void X()" in all_text
    # Header applied
    hdr_text = "\n".join(
        p.text for p in doc.sections[0].header.paragraphs
    )
    assert "CoreTax" in hdr_text
    assert "Solution Architect" in hdr_text
    # Table present
    assert len(doc.tables) >= 1
    first_row = [c.text for c in doc.tables[0].rows[0].cells]
    assert "Name" in first_row and "Layer" in first_row


def test_figure_skips_missing_file(tmp_path: Path):
    b = DocxBuilder("T", "S", "Tenant", "Persona")
    # Should not raise even though the file doesn't exist.
    b.figure(str(tmp_path / "missing.png"), "caption")
    out = b.save(str(tmp_path / "f.docx"))
    assert Path(out).exists()
