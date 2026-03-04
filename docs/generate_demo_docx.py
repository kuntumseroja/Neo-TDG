"""Generate the Neo-TDG Demo Scenario document in DOCX format."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    s = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(s)


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1B4F72")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EBF5FB")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_heading_with_number(doc, number, text, level=1):
    """Add a numbered heading."""
    doc.add_heading(f"{number}. {text}", level=level)


def main():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ============================================================
    # TITLE PAGE
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Neo-TDG")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(27, 79, 114)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("SDLC Knowledge Engine")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(52, 73, 94)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Demo Scenario & Narration Guide")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("RAG-Powered Technical Documentation & Code Intelligence for CoreTax")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(120, 120, 120)

    for _ in range(4):
        doc.add_paragraph()

    footer_info = doc.add_paragraph()
    footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_info.add_run("Prepared for: DJP CoreTax Development Team")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ============================================================
    # TABLE OF CONTENTS
    # ============================================================
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Demo Overview",
        "2. Pre-Demo Setup",
        "3. Scenario 1: Application Initialization",
        "4. Scenario 2: Solution Crawler",
        "5. Scenario 3: Document Generation & Download",
        "6. Scenario 4: Knowledge Store Ingestion",
        "7. Scenario 5: RAG Knowledge Chat",
        "8. Scenario 6: Flow Explorer",
        "9. Scenario 7: SDLC Tools - Bug Resolution",
        "10. Scenario 8: SDLC Tools - Test Generator",
        "11. Scenario 9: SDLC Tools - Architecture Validator",
        "12. Scenario 10: End-to-End Workflow",
        "13. Feature Summary Matrix",
        "Appendix A: Sample RAG Chat Questions",
        "Appendix B: CoreTaxSample Architecture",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ============================================================
    # 1. DEMO OVERVIEW
    # ============================================================
    add_heading_with_number(doc, 1, "Demo Overview")

    doc.add_heading("Purpose", level=2)
    doc.add_paragraph(
        "Demonstrate how Neo-TDG accelerates the Software Development Life Cycle (SDLC) "
        "for the CoreTax system by providing AI-powered code intelligence, automated "
        "documentation, and developer productivity tools."
    )

    doc.add_heading("Target Audience", level=2)
    audiences = [
        "Development Team Leads",
        "Software Architects",
        "Project Managers",
        "QA Engineers",
        "DJP Technical Stakeholders",
    ]
    for a in audiences:
        doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("Key Value Propositions", level=2)
    add_table(doc,
        ["Value", "Description"],
        [
            ["Automated Discovery", "Crawl .NET solutions to discover architecture, endpoints, consumers, schedulers, and integrations automatically"],
            ["AI-Powered Knowledge", "RAG-based chat understands your entire codebase and answers technical questions"],
            ["Documentation Generation", "Auto-generate comprehensive Markdown and PDF reports from crawl results"],
            ["Flow Tracing", "Trace business flows end-to-end across microservices"],
            ["SDLC Acceleration", "Bug resolution, test generation, and architecture validation tools"],
        ],
        col_widths=[4, 13],
    )

    doc.add_heading("Tech Stack Covered in Demo", level=2)
    add_table(doc,
        ["Technology", "Role"],
        [
            ["C# / .NET 8", "Backend Services"],
            ["Angular", "Frontend SPA"],
            ["CQRS with MediatR", "Command/Query Separation"],
            ["MassTransit + RabbitMQ", "Async Message Bus"],
            ["Redis", "Distributed Caching"],
            ["Hangfire", "Background Job Scheduling"],
            ["Consul", "Service Discovery"],
            ["ELK Stack", "Centralized Logging"],
            ["AWS S3", "Object/Document Storage"],
            ["Entity Framework Core", "ORM / Data Access"],
        ],
        col_widths=[5, 12],
    )

    doc.add_page_break()

    # ============================================================
    # 2. PRE-DEMO SETUP
    # ============================================================
    add_heading_with_number(doc, 2, "Pre-Demo Setup")

    doc.add_heading("Prerequisites", level=2)
    add_table(doc,
        ["Item", "Requirement"],
        [
            ["Ollama", "Running locally with deepseek-coder-v2:16b and nomic-embed-text models"],
            ["Python", "3.9+ with all dependencies (pip install -r requirements.txt)"],
            ["Browser", "Chrome (latest)"],
            ["Example Project", "examples/CoreTaxSample/ directory present"],
        ],
        col_widths=[4, 13],
    )

    doc.add_heading("Startup Steps", level=2)
    steps = [
        ("Start Ollama", "ollama serve"),
        ("Verify Models", "ollama list  (should show deepseek-coder-v2:16b, nomic-embed-text)"),
        ("Start Neo-TDG", "streamlit run app.py --server.port 8502"),
        ("Open Browser", "Navigate to http://localhost:8502"),
    ]
    for i, (title, cmd) in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Step {i}: {title}")
        run.bold = True
        run.font.size = Pt(10)
        p2 = doc.add_paragraph(cmd)
        for run in p2.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)

    doc.add_page_break()

    # ============================================================
    # 3. SCENARIO 1: APPLICATION INITIALIZATION
    # ============================================================
    add_heading_with_number(doc, 3, "Scenario 1: Application Initialization")

    doc.add_heading("Objective", level=2)
    doc.add_paragraph("Show how Neo-TDG connects to the AI backend and initializes all services.")

    doc.add_heading("Demo Steps", level=2)
    add_table(doc,
        ["Step", "Action", "Expected Result"],
        [
            ["1", "Open browser at http://localhost:8502", "Neo-TDG landing page loads"],
            ["2", "Observe the sidebar", "Settings panel with Ollama Configuration and Knowledge Store"],
            ["3", 'Expand "Ollama Configuration"', "Shows URL, LLM Model dropdown, Embedding Model"],
            ["4", "Select deepseek-coder-v2:16b from dropdown", "Model selection updates"],
            ["5", 'Click "Initialize / Reconnect"', "Service initialization begins"],
            ["6", "Observe Service Status indicators", "LLM turns GREEN, Store turns GREEN"],
        ],
        col_widths=[1.5, 6, 9.5],
    )

    doc.add_heading("Narration Script", level=2)
    narration = doc.add_paragraph()
    narration.style = "Intense Quote"
    narration.add_run(
        '"Welcome to Neo-TDG, the SDLC Knowledge Engine for CoreTax. This tool leverages '
        "AI-powered RAG technology to provide deep code intelligence across our entire solution.\n\n"
        "On the left sidebar, we configure the AI backend. We're using Ollama with the "
        "deepseek-coder model for code understanding, and nomic-embed-text for vector embeddings.\n\n"
        "When I click Initialize, the system connects to the LLM, sets up the vector knowledge "
        'store, and prepares all analysis engines. Notice both status indicators are now green '
        '-- we\'re ready to go."'
    )

    doc.add_page_break()

    # ============================================================
    # 4. SCENARIO 2: SOLUTION CRAWLER
    # ============================================================
    add_heading_with_number(doc, 4, "Scenario 2: Solution Crawler")

    doc.add_heading("Objective", level=2)
    doc.add_paragraph(
        "Demonstrate automated discovery of a .NET solution's architecture, endpoints, "
        "consumers, schedulers, integrations, and data models."
    )

    doc.add_heading("Demo Steps", level=2)
    add_table(doc,
        ["Step", "Action", "Expected Result"],
        [
            ["1", 'Click "Solution Crawler" in Navigation', "Solution Crawler page loads"],
            ["2", "Enter path to CoreTaxSample directory", "Path appears in input"],
            ["3", 'Click "Crawl Solution"', "Progress bar advances through each project"],
            ["4", "Wait for completion", "Success: 8 projects, 19 endpoints, 6 consumers, 8 schedulers"],
            ["5", 'Click "Projects" tab', "8 projects listed with layer classification"],
            ["6", "Scroll to Dependency Graph", "Mermaid diagram of project dependencies"],
            ["7", 'Click "Endpoints" tab', "19 API endpoints with HTTP method, route, auth status"],
            ["8", 'Click "Consumers" tab', "6 MassTransit consumers including InvoiceApprovalSaga"],
            ["9", 'Click "Schedulers" tab', "8 schedulers (Hangfire, BackgroundService, etc.)"],
            ["10", 'Click "Integrations" tab', "7 integrations: Consul, gRPC, HTTP, RabbitMQ, Redis, S3"],
            ["11", 'Click "Data Models" tab', "5 EF Core entities from CoreTaxDbContext"],
        ],
        col_widths=[1.5, 6, 9.5],
    )

    doc.add_heading("Narration Script", level=2)
    narration = doc.add_paragraph()
    narration.style = "Intense Quote"
    narration.add_run(
        "\"Now let's see the Solution Crawler in action. I'll point it to our CoreTaxSample "
        "solution -- a realistic .NET 8 solution simulating the Indonesian tax system.\n\n"
        "The crawler parses the .sln file, discovers projects, and analyzes each one. "
        "It reads controllers for endpoints, scans for MassTransit consumers, finds Hangfire "
        "jobs, and detects integrations with Redis, RabbitMQ, Consul, and more.\n\n"
        "Results: 8 projects, 19 API endpoints, 6 message consumers, and 8 scheduled jobs "
        "discovered automatically.\n\n"
        "In Projects, we see Clean Architecture layers. The dependency graph shows how "
        "Presentation depends on Application, which depends on Domain.\n\n"
        "Endpoints reveals all REST APIs with Auth indicators. Consumers found MassTransit "
        "event consumers including the InvoiceApprovalSaga state machine.\n\n"
        "Schedulers show Hangfire cron jobs and background services. Integrations are grouped "
        "by type -- Redis, RabbitMQ, Consul, S3, HTTP, and gRPC.\n\n"
        'Data Models shows all 5 entities managed by CoreTaxDbContext."'
    )

    doc.add_heading("Crawl Results Summary", level=2)
    add_table(doc,
        ["Category", "Count", "Details"],
        [
            ["Projects", "8", "Domain, Application, Infrastructure, Presentation, Worker, Contracts, Shared, Angular"],
            ["Endpoints", "19", "REST APIs across 6 controllers with Auth detection"],
            ["Consumers", "6", "5 IConsumer<T> + 1 MassTransit Saga"],
            ["Schedulers", "8", "3 Hangfire Recurring + 1 Enqueue + 1 Schedule + 2 BackgroundService + 1 HostedService"],
            ["Integrations", "7", "Consul, gRPC, HTTP (x2), RabbitMQ, Redis, S3"],
            ["Data Models", "5", "Taxpayer, TaxInvoice, TaxReturn, TaxPayment, TaxAudit"],
        ],
        col_widths=[3, 2, 12],
    )

    doc.add_page_break()

    # ============================================================
    # 5. SCENARIO 3: DOCUMENT GENERATION
    # ============================================================
    add_heading_with_number(doc, 5, "Scenario 3: Document Generation & Download")

    doc.add_heading("Objective", level=2)
    doc.add_paragraph("Show automatic documentation generation in Markdown and PDF formats with download capability.")

    doc.add_heading("Demo Steps", level=2)
    add_table(doc,
        ["Step", "Action", "Expected Result"],
        [
            ["1", "After crawl, scroll to download section", 'Two buttons: "Download MD Report" and "Download PDF Report"'],
            ["2", 'Click "Download MD Report"', "Browser downloads CoreTaxSample_report.md"],
            ["3", "Open the MD file", "Full technical documentation with tables and diagrams"],
            ["4", 'Click "Download PDF Report"', "Browser downloads CoreTaxSample_report.pdf"],
            ["5", "Open the PDF file", "Formatted PDF with headings, tables, code blocks"],
        ],
        col_widths=[1.5, 6, 9.5],
    )

    doc.add_heading("Narration Script", level=2)
    narration = doc.add_paragraph()
    narration.style = "Intense Quote"
    narration.add_run(
        "\"After every crawl, Neo-TDG generates comprehensive technical documentation in two formats.\n\n"
        "The Markdown report contains everything -- solution overview, architecture diagram, "
        "project details, API endpoints, consumers, schedulers, integration points, and data models.\n\n"
        "The PDF version provides the same in a print-ready format for stakeholders.\n\n"
        'This eliminates hours of manual documentation work. Every crawl produces instantly up-to-date docs."'
    )

    doc.add_heading("Generated Report Sections", level=2)
    sections = [
        "Title & Metadata (solution name, crawl timestamp)",
        "Solution Overview (metrics table)",
        "Layer Distribution (projects per architecture layer)",
        "Architecture Diagram (Mermaid graph)",
        "Project Details (layer, framework, references, NuGet packages)",
        "API Endpoints (method, route, controller, auth)",
        "Message Consumers (class, message type, queue)",
        "Scheduled Jobs (name, cron, description)",
        "Integration Points (grouped by type)",
        "Data Models (entity, DbContext, properties)",
        "Dependency Graph (Mermaid diagram)",
    ]
    for i, s in enumerate(sections, 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.add_page_break()

    # ============================================================
    # 6. SCENARIO 4: KNOWLEDGE STORE INGESTION
    # ============================================================
    add_heading_with_number(doc, 6, "Scenario 4: Knowledge Store Ingestion")

    doc.add_heading("Objective", level=2)
    doc.add_paragraph("Show how crawl results are ingested into the RAG knowledge base for AI-powered querying.")

    doc.add_heading("Demo Steps", level=2)
    add_table(doc,
        ["Step", "Action", "Expected Result"],
        [
            ["1", 'Click "Ingest to Knowledge Store"', 'Spinner: "Ingesting crawl report..."'],
            ["2", "Wait for completion", '"Ingested crawl report: X chunks created"'],
            ["3", 'Navigate to "Knowledge Store"', "Knowledge Store page loads"],
            ["4", "Observe Statistics", "Total chunks, documents, indexed services"],
            ["5", "Observe Chunk Type Distribution", "Counts per type: overview, architecture, component, etc."],
            ["6", "Scroll to Document Browser", "Lists all ingested documents"],
        ],
        col_widths=[1.5, 6, 9.5],
    )

    doc.add_heading("Narration Script", level=2)
    narration = doc.add_paragraph()
    narration.style = "Intense Quote"
    narration.add_run(
        "\"Now let's make this crawl data queryable by AI. Clicking 'Ingest to Knowledge Store' "
        "converts results into vector-embedded chunks that our RAG engine can search.\n\n"
        "The system creates structured chunks for each aspect -- endpoints, consumers, and "
        "architecture become searchable knowledge.\n\n"
        'This is the foundation for AI-powered code intelligence -- the knowledge store is '
        'our codebase\'s memory."'
    )

    doc.add_heading("Knowledge Store Capabilities", level=2)
    add_table(doc,
        ["Feature", "Description"],
        [
            ["Markdown Ingestion", "Import individual .md files or entire directories"],
            ["TechDocGen Output", "Import pre-generated TechDocGen documentation"],
            ["Full Rebuild", "Drop and rebuild entire index from scratch"],
            ["Document Browser", "View and delete individual documents"],
            ["Vector Search", "Cosine similarity with ChromaDB"],
            ["Hybrid Search", "BM25 + vector fusion reranking"],
            ["Metadata Filtering", "Filter by service, domain, chunk type"],
        ],
        col_widths=[5, 12],
    )

    doc.add_page_break()

    # ============================================================
    # 7. SCENARIO 5: RAG KNOWLEDGE CHAT
    # ============================================================
    add_heading_with_number(doc, 7, "Scenario 5: RAG Knowledge Chat")

    doc.add_heading("Objective", level=2)
    doc.add_paragraph("Demonstrate AI-powered conversational querying of the codebase using five analysis modes.")

    modes = [
        ("5A: Explain Mode", "explain",
         "How does the tax invoice submission process work?",
         "Detailed explanation with component references, flow description, source citations"),
        ("5B: Find Mode", "find",
         "Where are the Redis caching implementations?",
         "Lists file paths, class names, and relevant code patterns"),
        ("5C: Trace Mode", "trace",
         "Trace the payment processing flow from API to completion",
         "Step-by-step flow from controller through handler, domain, event, consumer"),
        ("5D: Impact Mode", "impact",
         "What would be affected if I change the Taxpayer entity?",
         "Lists downstream handlers, DTOs, controllers, consumers, tests impacted"),
        ("5E: Test Mode", "test",
         "Suggest test cases for InvoiceSubmittedConsumer",
         "Unit tests, integration tests, edge cases with C# xUnit skeletons"),
    ]

    for mode_name, mode_val, question, expected in modes:
        doc.add_heading(f"Sub-Scenario {mode_name}", level=2)
        add_table(doc,
            ["Step", "Action", "Expected Result"],
            [
                ["1", f'Select Query Mode: "{mode_val}"', f"{mode_val.capitalize()} mode selected"],
                ["2", f'Type: "{question}"', "Question submitted"],
                ["3", "Observe response", expected],
            ],
            col_widths=[1.5, 7, 8.5],
        )

    doc.add_heading("Narration Script", level=2)
    narration = doc.add_paragraph()
    narration.style = "Intense Quote"
    narration.add_run(
        "\"This is where Neo-TDG truly shines -- the RAG Knowledge Chat understands our entire "
        "codebase in five specialized modes.\n\n"
        "EXPLAIN mode provides comprehensive answers referencing actual components with source citations.\n\n"
        "FIND mode is smart code search -- pinpointing exact classes and files, not just keywords.\n\n"
        "TRACE mode traces flows step by step -- from API controller through handlers, domain "
        "logic, events, and consumers.\n\n"
        "IMPACT mode helps change planning by identifying every downstream dependency.\n\n"
        'TEST mode generates actual test code with unit tests, integration tests, and edge cases."'
    )

    doc.add_heading("Chat Features", level=2)
    add_table(doc,
        ["Feature", "Description"],
        [
            ["Conversation History", "Persistent conversations with browse and reload"],
            ["Metadata Filters", "Filter by Service Name, Probis Domain, Chunk Type"],
            ["Source Citations", "Every answer shows source files with relevance scores"],
            ["Confidence Badges", "High (green), Medium (orange), Low (red) indicators"],
            ["Mermaid Diagrams", "Sequence and flow diagrams rendered inline"],
            ["Related Topics", "Suggested follow-up questions"],
        ],
        col_widths=[5, 12],
    )

    doc.add_page_break()

    # ============================================================
    # 8. SCENARIO 6: FLOW EXPLORER
    # ============================================================
    add_heading_with_number(doc, 8, "Scenario 6: Flow Explorer")

    doc.add_heading("Objective", level=2)
    doc.add_paragraph("Demonstrate visual flow tracing and component explanation capabilities.")

    doc.add_heading("Sub-Scenario 6A: Trace Flow", level=2)
    add_table(doc,
        ["Step", "Action", "Expected Result"],
        [
            ["1", 'Click "Flow Explorer" in Navigation', "Flow Explorer loads with two tabs"],
            ["2", 'Select "Trace Flow" tab', "Flow tracing interface appears"],
            ["3", "Enter: POST /api/v1/invoices/submit", "Entry point entered"],
            ["4", 'Click "Trace Flow"', "AI generates flow analysis"],
            ["5", "Observe Sequence Diagram", "Mermaid diagram: Controller -> Handler -> Domain -> Repo -> Event -> Consumer"],
            ["6", "Observe Step-by-Step breakdown", "Each step with icon, component, action, file reference"],
        ],
        col_widths=[1.5, 6, 9.5],
    )

    doc.add_heading("Sub-Scenario 6B: Explain Component", level=2)
    add_table(doc,
        ["Step", "Action", "Expected Result"],
        [
            ["1", 'Select "Explain Component" tab', "Component explanation interface"],
            ["2", "Enter file path to a handler", "Path entered"],
            ["3", 'Select Explain Type: "Class"', "Class mode selected"],
            ["4", 'Click "Explain"', "AI generates explanation"],
            ["5", "Observe output", "Purpose, dependencies, events, business rules"],
        ],
        col_widths=[1.5, 6, 9.5],
    )

    doc.add_heading("Flow Step Types", level=2)
    add_table(doc,
        ["Icon", "Type", "Description"],
        [
            ["->", "HTTP Entry", "Controller endpoint receiving the request"],
            [">>", "Command", "CQRS command dispatched"],
            ["**", "Handler", "Command/Query handler processing"],
            ["~~", "Domain Logic", "Business rule execution"],
            ["[]", "Repository", "Database operation"],
            ["<>", "Event", "Domain event published"],
            ["<<", "Consumer", "Async event consumer processing"],
        ],
        col_widths=[2, 4, 11],
    )

    doc.add_heading("Narration Script", level=2)
    narration = doc.add_paragraph()
    narration.style = "Intense Quote"
    narration.add_run(
        "\"The Flow Explorer provides visual tracing. The sequence diagram shows the full flow -- "
        "HTTP request, command dispatch, handler processing, persistence, event publishing, "
        "and async consumer processing.\n\n"
        "Each step includes actual file and line references.\n\n"
        "The Component Explainer deep-dives into any class -- its purpose, constructor dependencies, "
        'domain events, and business rules."'
    )

    doc.add_page_break()

    # ============================================================
    # 9. SCENARIO 7: BUG RESOLUTION
    # ============================================================
    add_heading_with_number(doc, 9, "Scenario 7: SDLC Tools - Bug Resolution")

    doc.add_heading("Objective", level=2)
    doc.add_paragraph("Show AI-powered bug analysis and resolution suggestions.")

    doc.add_heading("Demo Steps", level=2)
    add_table(doc,
        ["Step", "Action", "Expected Result"],
        [
            ["1", 'Navigate to "SDLC Tools"', "SDLC Tools page with 3 tabs"],
            ["2", 'Select "Bug Resolution" tab', "Bug resolution interface"],
            ["3", "Enter bug description", "Description entered (see sample below)"],
            ["4", "Enter stack trace (optional)", "Stack trace pasted"],
            ["5", 'Click "Analyze Bug"', "AI analyzes against codebase knowledge"],
            ["6", "Observe Summary", "Bug summary with severity: HIGH"],
            ["7", "Observe Affected Components", "FileTaxReturnHandler, TaxReturn, TaxReturnController"],
            ["8", "Observe Probable Causes", "Ranked with confidence scores (e.g., 0.85)"],
            ["9", "Observe Suggested Fixes", "Code changes with risk level"],
            ["10", "Expand Verification Tests", "C# xUnit test code for the fix"],
        ],
        col_widths=[1.5, 5, 10.5],
    )

    doc.add_heading("Sample Bug Description", level=2)
    p = doc.add_paragraph(
        "Tax return filing fails with NullReferenceException when taxpayer has no registered "
        "address. The SPT submission endpoint returns 500 error."
    )
    for run in p.runs:
        run.italic = True

    doc.add_heading("Narration Script", level=2)
    narration = doc.add_paragraph()
    narration.style = "Intense Quote"
    narration.add_run(
        "\"The Bug Resolution Assistant analyzes bugs against codebase knowledge. "
        "It identifies severity, affected components, probable causes ranked by confidence, "
        "and suggested fixes with risk levels.\n\n"
        "It even provides verification test cases in C# that we can copy directly into our "
        'test project. This turns hours-long debugging into minutes."'
    )

    doc.add_heading("Analysis Output Structure", level=2)
    add_table(doc,
        ["Section", "Content"],
        [
            ["Summary", "Concise bug description"],
            ["Severity", "low / medium / high / critical"],
            ["Affected Components", "List of impacted classes/modules"],
            ["Probable Causes", "Description, confidence score (0-1), component, file location"],
            ["Suggested Fixes", "Fix description, code location, code change, risk level"],
            ["Verification Tests", "C# xUnit test name, description, code skeleton"],
        ],
        col_widths=[5, 12],
    )

    doc.add_page_break()

    # ============================================================
    # 10. SCENARIO 8: TEST GENERATOR
    # ============================================================
    add_heading_with_number(doc, 10, "Scenario 8: SDLC Tools - Test Generator")

    doc.add_heading("Objective", level=2)
    doc.add_paragraph("Demonstrate automated test case generation for components.")

    doc.add_heading("Demo Steps", level=2)
    add_table(doc,
        ["Step", "Action", "Expected Result"],
        [
            ["1", 'Select "Test Generator" tab', "Test generation interface"],
            ["2", "Enter component path (e.g., SubmitInvoiceHandler.cs)", "Path entered"],
            ["3", 'Select "Unit Tests"', "Unit test mode"],
            ["4", 'Click "Generate Tests"', "C# xUnit test class generated"],
            ["5", "Observe output", "Arrange-Act-Assert pattern with mocked dependencies"],
            ["6", 'Select "Edge Cases"', "Edge case mode"],
            ["7", 'Click "Generate Tests"', "Edge case scenarios generated"],
            ["8", "Observe output", "Scenario name, description, input, expected behavior"],
        ],
        col_widths=[1.5, 7, 8.5],
    )

    doc.add_heading("Test Types", level=2)
    add_table(doc,
        ["Type", "Output Format"],
        [
            ["Unit Tests", "xUnit C# test class with mocked deps, Arrange-Act-Assert pattern"],
            ["Integration Tests", "Multi-component test with database/service setup"],
            ["Edge Cases", "Scenario name, description, input scenario, expected behavior"],
        ],
        col_widths=[5, 12],
    )

    doc.add_heading("Narration Script", level=2)
    narration = doc.add_paragraph()
    narration.style = "Intense Quote"
    narration.add_run(
        "\"The Test Generator creates test skeletons for any component. Unit Tests follow "
        "Arrange-Act-Assert with proper mocking. Edge Cases identify scenarios we might miss -- "
        'zero line items, max threshold, concurrent submissions, deactivated taxpayers."'
    )

    doc.add_page_break()

    # ============================================================
    # 11. SCENARIO 9: ARCHITECTURE VALIDATOR
    # ============================================================
    add_heading_with_number(doc, 11, "Scenario 9: SDLC Tools - Architecture Validator")

    doc.add_heading("Objective", level=2)
    doc.add_paragraph("Show automated validation of architecture rules against the crawled solution.")

    doc.add_heading("Demo Steps", level=2)
    add_table(doc,
        ["Step", "Action", "Expected Result"],
        [
            ["1", 'Select "Architecture Validator" tab', "Validation interface"],
            ["2", "Confirm rules file (coretax_rules.yaml)", "Path shown"],
            ["3", 'Click "Validate Architecture"', "Runs against last crawl report"],
            ["4", "Observe Metrics", "Total rules, Passed, Failed, Warnings"],
            ["5", "Observe Violations", "Rule name, severity, file, description, suggested fix"],
        ],
        col_widths=[1.5, 6, 9.5],
    )

    doc.add_heading("Severity Levels", level=2)
    add_table(doc,
        ["Icon", "Severity", "Meaning"],
        [
            ["!!!", "Error", "Critical architecture violation -- must fix"],
            ["!", "Warning", "Best practice deviation -- should address"],
            ["i", "Info", "Recommendation for improvement"],
        ],
        col_widths=[2, 4, 11],
    )

    doc.add_heading("Narration Script", level=2)
    narration = doc.add_paragraph()
    narration.style = "Intense Quote"
    narration.add_run(
        "\"The Architecture Validator enforces design rules automatically against a YAML rule set. "
        "It catches violations like Presentation directly referencing Infrastructure and "
        'suggests the proper dependency path. This ensures Clean Architecture is maintained."'
    )

    doc.add_page_break()

    # ============================================================
    # 12. SCENARIO 10: END-TO-END WORKFLOW
    # ============================================================
    add_heading_with_number(doc, 12, "Scenario 10: End-to-End Workflow")

    doc.add_heading("Complete Demo Timeline", level=2)
    add_table(doc,
        ["Step", "Action", "Duration"],
        [
            ["1", "Initialize services (LLM + Knowledge Store)", "~5 sec"],
            ["2", "Crawl CoreTaxSample solution", "~10 sec"],
            ["3", "Review crawl results across all tabs", "~2 min"],
            ["4", "Download MD and PDF reports", "~5 sec"],
            ["5", "Ingest crawl report to Knowledge Store", "~10 sec"],
            ["6", 'RAG Chat: "Explain the CQRS pattern"', "~15 sec"],
            ["7", 'RAG Chat: "Trace taxpayer registration flow"', "~15 sec"],
            ["8", "Flow Explorer: trace an endpoint", "~15 sec"],
            ["9", "Bug Resolution: analyze sample bug", "~15 sec"],
            ["10", "Test Generator: generate tests for handler", "~15 sec"],
            ["", "TOTAL DEMO TIME", "~5 minutes"],
        ],
        col_widths=[1.5, 9, 6.5],
    )

    doc.add_heading("Closing Narration", level=2)
    narration = doc.add_paragraph()
    narration.style = "Intense Quote"
    narration.add_run(
        "\"In just five minutes, we demonstrated the full power of Neo-TDG:\n\n"
        "- Crawled an 8-project .NET solution discovering 19 endpoints, 6 consumers, 8 schedulers, "
        "7 integrations, and 5 data models\n"
        "- Generated complete technical documentation in Markdown and PDF\n"
        "- Ingested everything into a RAG knowledge store for AI querying\n"
        "- Asked natural language questions about architecture, flows, and impact\n"
        "- Used SDLC tools for bug resolution, test generation, and architecture validation\n\n"
        "Neo-TDG transforms how teams understand, document, and maintain the CoreTax system. "
        'It\'s not just a tool -- it\'s a knowledge engine that grows with your codebase."'
    )

    doc.add_page_break()

    # ============================================================
    # 13. FEATURE SUMMARY MATRIX
    # ============================================================
    add_heading_with_number(doc, 13, "Feature Summary Matrix")

    add_table(doc,
        ["Feature", "Page", "Key Benefit"],
        [
            ["Solution Crawler", "Solution Crawler", "Eliminate manual architecture documentation"],
            ["Project Analysis", "Crawler > Projects", "Understand solution structure instantly"],
            ["Endpoint Discovery", "Crawler > Endpoints", "Complete API inventory"],
            ["Consumer Discovery", "Crawler > Consumers", "Map async message flows"],
            ["Scheduler Discovery", "Crawler > Schedulers", "Inventory automated processes"],
            ["Integration Detection", "Crawler > Integrations", "Map external dependencies"],
            ["Data Model Discovery", "Crawler > Data Models", "Understand data architecture"],
            ["Dependency Graph", "Crawler > Projects", "See architecture at a glance"],
            ["MD Report", "Solution Crawler", "Instant technical docs"],
            ["PDF Report", "Solution Crawler", "Shareable formatted reports"],
            ["Knowledge Ingestion", "Knowledge Store", "Enable AI-powered queries"],
            ["RAG Chat - Explain", "RAG Chat", "Deep code understanding"],
            ["RAG Chat - Find", "RAG Chat", "Find anything quickly"],
            ["RAG Chat - Trace", "RAG Chat", "Understand complex flows"],
            ["RAG Chat - Impact", "RAG Chat", "Safe refactoring"],
            ["RAG Chat - Test", "RAG Chat", "Better test coverage"],
            ["Flow Tracing", "Flow Explorer", "Visual sequence diagrams"],
            ["Component Explainer", "Flow Explorer", "Onboard developers faster"],
            ["Bug Resolution", "SDLC Tools", "Debug faster"],
            ["Test Generator", "SDLC Tools", "Accelerate QA"],
            ["Architecture Validator", "SDLC Tools", "Enforce design standards"],
        ],
        col_widths=[5, 5, 7],
    )

    doc.add_page_break()

    # ============================================================
    # APPENDIX A: SAMPLE QUESTIONS
    # ============================================================
    doc.add_heading("Appendix A: Sample RAG Chat Questions", level=1)

    question_groups = [
        ("Explain Mode", [
            "How does the CQRS pattern work in this solution?",
            "Explain the tax invoice lifecycle",
            "What is the role of CoreTax.Shared project?",
            "How does authentication work across controllers?",
        ]),
        ("Find Mode", [
            "Where are the Hangfire job definitions?",
            "Find all classes that use IDistributedCache",
            "Where is Consul service registration configured?",
            "Find the S3 document storage implementation",
        ]),
        ("Trace Mode", [
            "Trace the taxpayer registration flow from API to database",
            "Trace what happens when a tax payment is processed",
            "Trace the audit scheduling flow",
            "Trace the invoice approval saga",
        ]),
        ("Impact Mode", [
            "What would be affected if I change the TaxInvoice entity?",
            "Impact of modifying the payment gateway client",
            "What breaks if I remove the Redis caching service?",
            "Impact of changing the NPWP validation rules",
        ]),
        ("Test Mode", [
            "Suggest tests for ProcessPaymentCommandHandler",
            "What edge cases should I test for tax return filing?",
            "Generate integration tests for InvoiceController",
            "What should I test in the AuditScheduledConsumer?",
        ]),
    ]

    for group_name, questions in question_groups:
        doc.add_heading(group_name, level=2)
        for q in questions:
            doc.add_paragraph(q, style="List Bullet")

    doc.add_page_break()

    # ============================================================
    # APPENDIX B: ARCHITECTURE
    # ============================================================
    doc.add_heading("Appendix B: CoreTaxSample Architecture", level=1)

    arch_items = [
        ("CoreTaxSample.sln", "Main solution file with 8 projects"),
        ("CoreTax.Domain", "Entities (Taxpayer, TaxInvoice, TaxReturn, TaxPayment, TaxAudit), "
         "Value Objects (Npwp, Money), Enums (TaxStatus, AuditStatus)"),
        ("CoreTax.Contracts", "5 Commands (RegisterTaxpayer, SubmitInvoice, FileTaxReturn, "
         "ProcessPayment, ScheduleAudit) + 5 Events"),
        ("CoreTax.Application", "5 Command Handlers (MediatR), 3 Query Handlers, "
         "2 FluentValidation Validators, 3 DTOs"),
        ("CoreTax.Infrastructure", "CoreTaxDbContext (EF Core, 5 DbSets), RedisCacheService, "
         "ConsulRegistrationService, ElasticsearchLogging, PaymentGatewayClient, DocumentStorageService (S3)"),
        ("CoreTax.Presentation", "6 Controllers: TaxRegistration (3), Invoice (4), "
         "TaxReturn (3), Payment (3), Audit (4), Report (2) = 19 endpoints total"),
        ("CoreTax.Worker", "5 MassTransit IConsumer<T>, 1 InvoiceApprovalSaga, "
         "HangfireJobScheduler (3 recurring + enqueue + schedule), 2 BackgroundServices"),
        ("CoreTax.Shared", "JwtTokenService, AuthorizationHandler, SerilogConfiguration, "
         "ServiceCollectionExtensions"),
        ("CoreTax.Angular", "5 feature modules (registration, invoice, filing, payment, dashboard) "
         "with Component + Routes + Service each"),
    ]

    for name, desc in arch_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    # ============================================================
    # SAVE
    # ============================================================
    output_path = "/Users/priyo/Downloads/DJP/Neo-TDG/docs/DEMO_SCENARIO.docx"
    doc.save(output_path)
    print(f"DOCX saved to: {output_path}")


if __name__ == "__main__":
    main()
